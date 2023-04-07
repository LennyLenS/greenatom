# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Mm

def print_hi(groupDepartmentCenter, directorPositions, jobObjectives, directorName, subordinationDiagram, mainResponsibilities, education, experience, skills, name_file):


    # создание пустого документа
    doc = Document()
    # данные таблицы без названий колонок
    #groupDepartmentCenter = "1"
    #directorPositions = "2"
    #jobObjectives = ["3", "4"]
    #directorName = "5"
    #subordinationDiagram = ["6", "7", "8"]
    #mainResponsibilities = "9"
    #education = "10"
    #experience = "11"
    #skills = "12"
    items = (
        ("Группа, Отдел, Центр", groupDepartmentCenter),
        ("Генеральный директор/Заместитель Генерального директора", directorPositions),
    )

    items_2 = [
        "1. Цель должности"
    ] + jobObjectives
    print("111111")
    print(jobObjectives)
    # .add_run("Должностная инструкция")
    run_2 = doc.add_paragraph(
        "УТВЕРЖДАЮ\n" + directorPositions + "\n_______________/" + directorName + "/\n«____»_______________ 202__ г."
    )

    run_2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = doc.add_paragraph("").add_run("                         Должностная инструкция")

    # название шрифта
    run.font.name = "Arial"
    # размер шрифта
    run.font.size = Pt(16)
    run.font.bold = True

    run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # fmt = run.paragraph_format
    # Форматируем:
    # отступ слева в мм
    # fmt.first_line_indent = Mm(20)
    # отступ сверху в мм
    # fmt.space_before = Mm(20)
    # отступ снизу в мм
    # fmt.space_after = Mm(10)

    # добавляем таблицу с одной строкой
    # для заполнения названий колонок
    table = doc.add_table(1, len(items[0]))
    table.style = "Table Grid"
    # Получаем строку с колонками из добавленной таблицы
    head_cells = table.rows[0].cells
    # добавляем названия колонок
    for i, item in enumerate(["Должность", "Специалист"]):
        p = head_cells[i].paragraphs[0]
        # название колонки
        p.add_run(item).bold = True
        # выравниваем посередине
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # добавляем данные к существующей таблице
    for k in range(len(items)):
        # добавляем строку с ячейками к объекту таблицы
        cells = table.add_row().cells
        for i, item in enumerate(items[k]):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            # если последняя ячейка
            if i == 1:
                # изменим шрифт
                cells[i].paragraphs[0].runs[0].font.name = "Arial"
    table = doc.add_table(len(items_2), 1)
    table.style = "Table Grid"

    print(len(items_2))
    for k in range(len(items_2)):
        p = (table.rows[k].cells)[0].paragraphs[0]

        p.add_run(items_2[k]).bold = k % 2 == 0

    p = doc.add_paragraph("")

    nachalbnica = subordinationDiagram
    table_2 = doc.add_table(len(nachalbnica), 1)
    table_2.style = "Table Grid"
    for i, item in enumerate(nachalbnica):
        p_1 = (table_2.rows[i].cells)[0].paragraphs[0]
        table_2.rows[i].cells[0].width = 3000000
        p_1.add_run(item).bold = False
        p_1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("")

    table_3 = doc.add_table(1, 1)
    table_3.style = "Table Grid"
    p_2 = (table_3.rows[0].cells)[0].paragraphs[0]
    p_2.add_run("3. Основные обязанности").bold = True

    items_3 = (
        (
            "%5",
            "Участие во внедрении, тиражировании, технологическом развитии портальных и интернет-систем:\nнастройка систем в целях внедрения, тиражирования;\nтестирование обновлений программного обеспечения;\nактуализация эксплуатационной документации, разработка инструкций для пользователей.\nПодготовка оценки трудозатрат на создание новых информационных систем",
        ),
        ("%5", "Администрирование портальных и интернет-систем:\nэксплуатационное обслуживание, мониторинг;\nадминистрирование справочных данных;\nнастройка;\nподключение, настройка и изменение прав доступа пользователей."),
        ("%5", "Решение инцидентов, связанных с прикладными системами:\nрешение проблем, формирование обходных решений;\nконсультация пользователей;\nанализ и формирование требований по развитию системы;"),
        ("%85", mainResponsibilities),
    )

    table_4 = doc.add_table(1, 2)
    table_4.style = "Table Grid"
    p_3 = (table_4.rows[0].cells)[0].paragraphs[0]
    p_3.add_run("% необходимого времени").bold = True
    p_3 = (table_4.rows[0].cells)[1].paragraphs[0]
    p_3.add_run("Прямые обязанности").bold = True
    for k in range(len(items_3)):
        # добавляем строку с ячейками к объекту таблицы
        cells = table_4.add_row().cells
        for i, item in enumerate(items_3[k]):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            if i == 1:
                cells[i].width = 20000000
        # если последняя ячейка
        # изменим шрифт
        # cells[i].paragraphs[0].runs[0].font.name = "Arial"
    table_4 = doc.add_table(1, 2)
    table_4.style = "Table Grid"
    p_3 = (table_4.rows[0].cells)[0].paragraphs[0]
    p_3.add_run("% необходимого времени").bold = True
    p_3 = (table_4.rows[0].cells)[1].paragraphs[0]
    p_3.add_run("Совместные обязанности").bold = True
    table_4 = doc.add_table(1, 2)
    table_4.style = "Table Grid"

    table_5 = doc.add_table(1, 1)
    table_5.style = "Table Grid"
    p_4 = (table_5.rows[0].cells)[0].paragraphs[0]
    p_4.add_run("4 Уровень контактов/коммуникаций").bold = True

    items_3 = (
        ("Функциональный руководитель."),
        ("Административный руководитель, работники подразделения, работнки смежных подразделений в рамках исполнения должгомтных обязанностей."),
        ("Ключевые или конечные пользователи по направлению граппы.")
    )

    table_6 = doc.add_table(len(items_3), 1)
    table_6.style = "Table Grid"
    for i, item in enumerate(items_3):
        p_6 = (table_6.rows[i].cells)[0].paragraphs[0]
        p_6.add_run(item).bold = False
        p_6.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_6 = doc.add_table(1, 1)
    table_6.style = "Table Grid"
    p_6 = (table_6.rows[0].cells)[0].paragraphs[0]
    p_6.add_run("5. Требования должности:").bold = True

    worker = (
        ("Образование:", education),
        ("Опыт работы:", experience),
        ("Форма допуска к государственной тайне:", "В соответвии с номенклатурой должностей"),
        ("Знания, навыки, личностно-деловые качества(компетенции):", skills + "\nВладение инструментарием автоматизации тестирования\nОсновы информационной безопасности web-ресурсов\nОсновы современных систем управления базами данных\nСетевые протоколы и основы web-технологий\nЛичностные-деловые качества: ответственность, вниматьльность, быстрая обучаемость, коммуникабельность, исполнительность, нацеленность на результат"),
        ("Знание языков:", "Английский, чтение технической литературы в области ИТ, со словарем"),
        ("Знание программного обеспечения:", "MS Office (Excel, Word, Visio)\nОперационные системы Windows\nGitLab / Jira / Confluence / OpenProject / AzureDevops"),
    )
    table_7 = doc.add_table(1, 2)
    table_7.style = "Table Grid"
    p_7 = (table_7.rows[0].cells)[0].paragraphs[0]
    p_7.add_run(worker[0][0]).bold = True
    p_7 = (table_7.rows[0].cells)[1].paragraphs[0]
    p_7.add_run(worker[0][1]).bold = False
    for k in range(1, len(worker)):
        # добавляем строку с ячейками к объекту таблицы
        cells = table_7.add_row().cells
        for i, item in enumerate(worker[k]):
            # вставляем данные в ячейки
            cells[i].text = str(item)
            # если последняя ячейка
            if i == 0:
                # изменим шрифт
                cells[i].paragraphs[0].runs[0].font.bold = True

    table_8 = doc.add_table(3, 1)
    table_8.style = "Table Grid"
    p_8 = (table_8.rows[0].cells)[0].paragraphs[0]
    p_8.add_run("6. Согласовано:").bold = True
    p_8 = (table_8.rows[1].cells)[0].paragraphs[0]
    p_8.add_run(
        "\nНачальник управления ______________________________________/__________________/\n                                                                    (подпись)                                   (ФИО)\n\nНачальник отдела          ______________________________________/__________________/\n                                                                    (подпись)                                   (ФИО)\n"
    ).bold = False
    p_8 = (table_8.rows[2].cells)[0].paragraphs[0]
    p_8.add_run(
        "\nС должностной инструкцией ознакомлен и обязуюсь соблюдать:\n______________________________________/___________________________________________________\n(подпись работника, дата)                                        (ФИО)\n"
    ).bold = True

    run_1 = doc.add_paragraph("").add_run("\n\n\nРаботник___________________")

    # название шрифта
    run_1.font.name = "Arial"
    # размер шрифта
    run_1.font.size = Pt(13)
    run.font.bold = False

    doc.save(name_file)


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    print_hi('PyCharm')

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
